"""
DocBot v3 — Word SEQ field helpers (W4).

Inserts ``SEQ Figure`` / ``SEQ Table`` Word fields via raw OOXML so that
Word's ``Update Field`` command regenerates TOF/TOT automatically.

This replaces the old plain-text caption approach where captions were
never registered with Word's caption-tracking machinery.

Design
------
Word SEQ fields work via the following OOXML structure:

    <w:r><w:fldChar w:fldCharType="begin"/></w:r>
    <w:r><w:instrText> SEQ Figure \\* ARABIC </w:instrText></w:r>
    <w:r><w:fldChar w:fldCharType="separate"/></w:r>
    <w:r><w:t>1</w:t></w:r>         <!-- cached value -->
    <w:r><w:fldChar w:fldCharType="end"/></w:r>

For module-prefixed numbering (NCB style): the module number is inserted
as literal text before the SEQ field, e.g. "Figure 10-" + SEQ.

For continuous numbering (NCD style): a pure ``SEQ Figure \\* ARABIC``
field is used, producing 1, 2, 3 … across the whole document.

The ``\\r N`` switch resets the counter to N at module boundaries (used
only when the previous module's figure count must restart).

``add_caption`` adds a styled paragraph containing the full caption:
    "<Figure|Table> <field> <caption_text>"
"""

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def add_caption(
    doc,
    prefix: str,
    caption_text: str,
    style_cfg,
    *,
    seq_name: str = "Figure",
    module_num: int | None = None,
    reset_to: int | None = None,
    font_size_pt: float = 10.0,
) -> None:
    """
    Add a Word-compatible captioned paragraph with a SEQ field.

    Args:
        doc:           The python-docx Document object.
        prefix:        Display prefix, e.g. "Figure" or "Table".
        caption_text:  The text that follows the number, e.g. "Login Screen".
        style_cfg:     The StyleConfig or equivalent with color helpers.
        seq_name:      SEQ counter name (``"Figure"`` or ``"Table"``).
        module_num:    If provided, inserts "10-" before the SEQ field
                       (module-prefixed numbering for NCB).
        reset_to:      If provided, resets the counter to this value via
                       the ``\\r`` switch (used at module boundaries).
        font_size_pt:  Caption font size.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)

    # Attempt to apply the built-in "Caption" style so TOF/TOT works
    try:
        p.style = doc.styles["Caption"]
    except Exception:
        pass  # Style not in template — will still use field XML

    # Prefix run: "Figure " or "Table " (bold, colored)
    _add_run(p, f"{prefix} ", bold=True, size_pt=font_size_pt,
             color_hex=style_cfg.get_color("secondary"))

    # Module prefix: "10-" (bold)
    if module_num is not None:
        _add_run(p, f"{module_num}-", bold=True, size_pt=font_size_pt,
                 color_hex=style_cfg.get_color("secondary"))

    # SEQ field run
    _insert_seq_field(p, seq_name=seq_name, reset_to=reset_to,
                      size_pt=font_size_pt, color_hex=style_cfg.get_color("secondary"))

    # Caption text run: " Login Screen"
    _add_run(p, f" {caption_text}", bold=False, size_pt=font_size_pt,
             color_hex=style_cfg.get_color("muted"), italic=True)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _add_run(
    para,
    text: str,
    bold: bool = False,
    size_pt: float = 10.0,
    color_hex: str = "333333",
    italic: bool = False,
) -> None:
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    try:
        r, g, b = _hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    except Exception:
        pass


def _insert_seq_field(
    para,
    seq_name: str = "Figure",
    reset_to: int | None = None,
    size_pt: float = 10.0,
    color_hex: str = "828282",
) -> None:
    """
    Insert a SEQ field into *para* using raw OOXML.

    Generates:
        SEQ Figure \\* ARABIC [\\r N]
    """
    # Build instruction text
    instr = f" SEQ {seq_name} \\* ARABIC"
    if reset_to is not None:
        instr += f" \\r {reset_to}"

    # Common run properties
    def _make_rpr(para) -> etree._Element:
        rpr = OxmlElement("w:rPr")
        if size_pt:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(size_pt * 2)))
            rpr.append(sz)
            szcs = OxmlElement("w:szCs")
            szcs.set(qn("w:val"), str(int(size_pt * 2)))
            rpr.append(szcs)
        r_color = OxmlElement("w:color")
        r_color.set(qn("w:val"), color_hex.lstrip("#").upper())
        rpr.append(r_color)
        return rpr

    p_el = para._p

    # begin
    r_begin = OxmlElement("w:r")
    r_begin.append(_make_rpr(para))
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_begin)
    p_el.append(r_begin)

    # instr text
    r_instr = OxmlElement("w:r")
    r_instr.append(_make_rpr(para))
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r_instr.append(it)
    p_el.append(r_instr)

    # separate
    r_sep = OxmlElement("w:r")
    r_sep.append(_make_rpr(para))
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    p_el.append(r_sep)

    # cached value "1" (Word will update on Ctrl+A, F9)
    r_val = OxmlElement("w:r")
    r_val.append(_make_rpr(para))
    wt = OxmlElement("w:t")
    wt.text = "1"
    r_val.append(wt)
    p_el.append(r_val)

    # end
    r_end = OxmlElement("w:r")
    r_end.append(_make_rpr(para))
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    p_el.append(r_end)


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

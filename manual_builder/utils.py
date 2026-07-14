"""
Word document helper utilities.

Consolidates XML-level helpers used by all renderers.
"""

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH


def hex_to_rgb(hex_str: str) -> RGBColor:
    """Convert a hex color string (with or without '#') to docx RGBColor."""
    hex_str = str(hex_str).lstrip("#")
    if len(hex_str) != 6:
        return RGBColor(27, 54, 93)  # fallback dark slate
    return RGBColor(*(int(hex_str[i:i + 2], 16) for i in (0, 2, 4)))


def set_cell_background(cell, fill_hex: str):
    """Set cell background shading."""
    fill_hex = str(fill_hex).lstrip("#")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table, border_color_hex: str = "CCCCCC", sz: str = "4"):
    """Set custom borders on all cells in a table."""
    border_color_hex = str(border_color_hex).lstrip("#")
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color_hex)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 180, right: int = 180):
    """Set cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for margin, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{margin}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_left_border(cell, color_hex: str, size_sz: int = 36):
    """Set a thick left border on a cell (for callout styling)."""
    color_hex = str(color_hex).lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_sz))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color_hex)
    tcBorders.append(left)
    for border_name in ["top", "bottom", "right"]:
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "none")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    """Remove all borders from a table (useful for layout tables)."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_bottom_border(paragraph, color_hex: str = "CCCCCC", sz: int = 6):
    """Add a bottom border to a paragraph (divider line)."""
    color_hex = str(color_hex).lstrip("#")
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_field(run):
    """Append the PAGE number field XML to a run."""
    for tag_type, text in [("begin", None), ("instrText", "PAGE"), ("separate", None), ("end", None)]:
        if tag_type == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
            run._r.append(el)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), tag_type)
            run._r.append(fc)


def add_numpages_field(run):
    """Append the NUMPAGES field XML to a run."""
    for tag_type, text in [("begin", None), ("instrText", "NUMPAGES"), ("separate", None), ("end", None)]:
        if tag_type == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
            run._r.append(el)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), tag_type)
            run._r.append(fc)


def add_toc_field(paragraph, max_depth: int = 3):
    """Add a native Word TOC field to a paragraph."""
    p_element = paragraph._p
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f'TOC \\o "1-{max_depth}" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    p_element.append(fldChar1)
    p_element.append(instrText)
    p_element.append(fldChar2)

    r = OxmlElement("w:r")
    rText = OxmlElement("w:t")
    rText.text = "[Right-click here and select 'Update Field' to generate Table of Contents]"
    r.append(rText)
    p_element.append(r)
    p_element.append(fldChar3)


def add_body_paragraph(doc, text: str, font_name: str = "Calibri", size_pt: float = 11,
                       color_hex: str = "000000"):
    """Add a standard body text paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = font_name
    r.font.size = Pt(size_pt)
    r.font.color.rgb = hex_to_rgb(color_hex)
    return p


def add_styled_heading(doc, text: str, level: int, style_config, numbering=None):
    """
    Add a heading at the given level using style config for font/size/color.

    If numbering string is provided, prepends it (e.g., "10.1 Add Case").
    """
    h = doc.add_heading(level=min(level, 4))
    display_text = f"{numbering}\t{text}" if numbering else text
    run = h.add_run(display_text)
    hcfg = style_config.heading_config(level)
    run.font.name = style_config.heading_font
    run.font.size = Pt(hcfg["size_pt"])
    run.font.bold = hcfg["bold"]
    run.font.color.rgb = hex_to_rgb(hcfg["color"])
    h.paragraph_format.space_before = Pt(hcfg.get("before_pt", 18))
    h.paragraph_format.space_after = Pt(hcfg.get("after_pt", 8))
    h.paragraph_format.keep_with_next = True
    return h


def setup_header_footer(doc, style_config, manifest=None):
    """
    Set up professional header/footer on all sections.

    Header: logo (left) + document title (right)
    Footer: company name (left) + page number (right)
    """
    logo_path_str = style_config.logo.get("path", "")
    logo_path = Path(logo_path_str) if logo_path_str else None
    font_name = style_config.body_font

    if manifest:
        doc_title = f"{manifest.client_display_name} — {manifest.manual_title}"
        company = manifest.client_display_name
    else:
        doc_title = "User Manual"
        company = ""

    for section in doc.sections:
        section.different_first_page_header_footer = True

        # ── HEADER ──
        header = section.header
        header.is_linked_to_previous = False
        hdr_element = header._element
        for child in list(hdr_element):
            hdr_element.remove(child)

        hdr_p = OxmlElement("w:p")
        hdr_element.append(hdr_p)

        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        table.autofit = False
        remove_table_borders(table)
        table.columns[0].width = Inches(3.25)
        table.columns[1].width = Inches(3.25)

        cell_logo = table.cell(0, 0)
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if logo_path and logo_path.exists():
            try:
                logo_size = style_config.logo.get("size_cm", 2)
                p_logo.add_run().add_picture(str(logo_path), height=Cm(logo_size))
            except Exception:
                pass

        cell_title = table.cell(0, 1)
        p_title = cell_title.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_title = p_title.add_run(doc_title)
        run_title.font.name = font_name
        run_title.font.size = Pt(8.5)
        run_title.font.italic = True
        run_title.font.color.rgb = RGBColor(100, 100, 100)

        sep_p = header.add_paragraph()
        add_bottom_border(sep_p, color_hex="CCCCCC", sz=4)
        sep_p.paragraph_format.space_before = Pt(2)
        sep_p.paragraph_format.space_after = Pt(0)

        # ── FOOTER ──
        footer = section.footer
        footer.is_linked_to_previous = False
        ftr_element = footer._element
        for child in list(ftr_element):
            ftr_element.remove(child)

        ftr_p_xml = OxmlElement("w:p")
        ftr_element.append(ftr_p_xml)

        ftr_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        ftr_table.autofit = False
        remove_table_borders(ftr_table)
        ftr_table.columns[0].width = Inches(3.25)
        ftr_table.columns[1].width = Inches(3.25)

        p_left = ftr_table.cell(0, 0).paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_company = p_left.add_run(company)
        r_company.font.name = font_name
        r_company.font.size = Pt(8)
        r_company.font.color.rgb = RGBColor(130, 130, 130)

        p_right = ftr_table.cell(0, 1).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_pre = p_right.add_run("Page ")
        r_pre.font.name = font_name
        r_pre.font.size = Pt(8)
        r_pre.font.color.rgb = RGBColor(100, 100, 100)
        add_page_field(p_right.add_run())
        r_mid = p_right.add_run(" of ")
        r_mid.font.name = font_name
        r_mid.font.size = Pt(8)
        r_mid.font.color.rgb = RGBColor(100, 100, 100)
        add_numpages_field(p_right.add_run())

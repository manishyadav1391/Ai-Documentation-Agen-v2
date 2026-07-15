"""Module section renderer."""

import json
from pathlib import Path
from typing import Any, Dict
from docx.shared import Inches, Pt, RGBColor
from manual_builder.utils import add_styled_heading, add_body_paragraph, hex_to_rgb
from manual_builder.renderers.screen import render_screen
from docbot.clients.profile import ClientProfile


def render_module(doc, session_dir, style, numbering):
    """
    Renders a complete recorded module from session data folder.
    Reads module_meta.json or screen meta to build module headings and summaries.
    """
    module_meta_path = session_dir / "module_meta.json"
    
    module_name = ""
    module_num = 10  # Fallback module number
    intro = ""
    features = []
    screen_order = []

    # 1. Load module metadata if it exists
    if module_meta_path.exists():
        try:
            with module_meta_path.open("r", encoding="utf-8") as f:
                meta = json.load(f)
            module_name = meta.get("module_name", "")
            module_num = meta.get("module_number", module_num)
            intro = meta.get("intro", "")
            features = meta.get("features", [])
            screen_order = meta.get("screen_order", [])
        except Exception as e:
            print(f"[Warning] Failed to read module_meta.json: {e}")

    # 2. Derive defaults from session folder or first screen if missing
    if not module_name:
        first_meta_path = session_dir / "screen_1_meta.json"
        if first_meta_path.exists():
            try:
                with first_meta_path.open("r", encoding="utf-8") as f:
                    first_meta = json.load(f)
                module_name = first_meta.get("screen_name") or first_meta.get("h1_text") or first_meta.get("title", "")
            except Exception:
                pass
        if not module_name:
            module_name = session_dir.name.replace("session_", "Session ").replace("_", " ").title()

    # Track current module numbering
    if numbering.mode == "continuous":
        # Increment level 1 section number dynamically
        current_section = numbering.enter_section(level=1)
        try:
            module_num = int(current_section.split(".")[0])
        except ValueError:
            module_num = 1
        numbering.set_current_module(module_num)
    else:
        numbering.set_current_module(module_num)
        numbering.set_section_number(level=1, number=module_num)

    # 3. Detect screens sequentially to check for single_screen layout
    if not screen_order:
        screen_files = sorted(session_dir.glob("screen_*_content.json"))
        for f in screen_files:
            try:
                idx = int(f.name.split("_")[1])
                screen_order.append(idx)
            except Exception:
                pass
        screen_order = sorted(list(set(screen_order)))

    screens_to_render = []
    for screen_idx in screen_order:
        content_path = session_dir / f"screen_{screen_idx}_content.json"
        meta_path = session_dir / f"screen_{screen_idx}_meta.json"
        
        if not content_path.exists() or not meta_path.exists():
            continue

        try:
            with meta_path.open("r", encoding="utf-8") as f:
                screen_meta = json.load(f)
            if screen_meta.get("state_of") is not None:
                continue
            screens_to_render.append((screen_idx, content_path, meta_path))
        except Exception:
            pass

    # Check for single-screen module layout under continuous mode
    single_screen = (len(screens_to_render) == 1) and (numbering.mode == "continuous")

    # Load client profile for voice rules / notes_block
    profile = None
    try:
        from config import get_config
        cfg = get_config()
        profile = ClientProfile.load(cfg.current_client)
    except Exception:
        pass

    if not single_screen:
        # Add Module Heading (e.g., "10. Case Form")
        heading_str = f"{module_num}."
        add_styled_heading(doc, module_name, level=1, style_config=style, numbering=heading_str)

        # 4. Intro Paragraph
        if intro:
            add_body_paragraph(doc, intro, font_name=style.body_font, size_pt=style.body_size, color_hex=style.get_color("body_text"))

        # Render voice notes_block if any
        if profile and profile.notes_block:
            p_note = doc.add_paragraph()
            p_note.paragraph_format.space_before = Pt(4)
            p_note.paragraph_format.space_after = Pt(8)
            r_note_lbl = p_note.add_run("Note: ")
            r_note_lbl.font.name = style.body_font
            r_note_lbl.font.bold = True
            r_note_lbl.font.size = Pt(style.body_size)
            r_note_lbl.font.color.rgb = hex_to_rgb(style.get_color("tertiary"))
            
            r_note_val = p_note.add_run(profile.notes_block)
            r_note_val.font.name = style.body_font
            r_note_val.font.size = Pt(style.body_size)
            r_note_val.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

        # 5. Features list
        if features:
            p_feat_hdr = doc.add_paragraph()
            p_feat_hdr.paragraph_format.space_before = Pt(6)
            p_feat_hdr.paragraph_format.space_after = Pt(4)
            r_feat_hdr = p_feat_hdr.add_run(f"The {module_name} module includes features for:")
            r_feat_hdr.font.name = style.body_font
            r_feat_hdr.font.bold = True
            r_feat_hdr.font.size = Pt(style.body_size)
            r_feat_hdr.font.color.rgb = hex_to_rgb(style.get_color("secondary"))

            for feat in features:
                p_feat = doc.add_paragraph(style="List Bullet")
                p_feat.paragraph_format.space_after = Pt(2)
                r_feat = p_feat.add_run(feat)
                r_feat.font.name = style.body_font
                r_feat.font.size = Pt(style.body_size)
                r_feat.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

    # 6. Render Screens sequentially
    for screen_idx, content_path, meta_path in screens_to_render:
        try:
            with content_path.open("r", encoding="utf-8") as f:
                content_data = json.load(f)
            with meta_path.open("r", encoding="utf-8") as f:
                screen_meta = json.load(f)
        except Exception as e:
            print(f"[Warning] Failed to load data for Screen {screen_idx}: {e}")
            continue

        print(f"  Rendering screen: {content_data.get('screen_name', f'Screen {screen_idx}')}...")

        if single_screen:
            # Render at level 1 directly (Heading 1) instead of level 2 (Heading 2)
            # Temporarily override enter_section level-2 behavior by making enter_section(2) call enter_section(1)
            # Or we can pass a parameter to render_screen or adjust NumberingTracker section counters.
            # In screen.py, it does: screen_number_str = numbering.enter_section(level=2)
            # To render at level 1, we can adjust enter_section or set level=1.
            # Wait, screen.py does screen_number_str = numbering.enter_section(level=2)
            # If we detect single_screen, we can make numbering.enter_section(level=2) return the level 1 section number!
            # Let's adjust enter_section in NumberingTracker or add a wrapper.
            # Wrapper:
            orig_enter_section = numbering.enter_section
            def single_screen_enter(level):
                if level == 2:
                    # Return current level 1 section number
                    return numbering.get_current_section_number()
                return orig_enter_section(level)
            numbering.enter_section = single_screen_enter

        render_screen(
            doc=doc,
            screen_index=screen_idx,
            session_dir=session_dir,
            content_data=content_data,
            screen_meta=screen_meta,
            style=style,
            numbering=numbering
        )

        if single_screen:
            # Restore enter_section
            numbering.enter_section = orig_enter_section

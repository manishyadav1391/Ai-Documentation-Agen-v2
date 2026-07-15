"""Field table (4-column) renderer."""

from typing import Any, Dict, List
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from manual_builder.utils import set_cell_background, set_cell_margins, set_table_borders, hex_to_rgb


def render_field_table(doc, screen_index, field_details, style, table_number_str):
    """
    Renders the 4-column field details table:
    Field Name | Utility | Information | Sample
    """
    table_prefix = style.numbering.get("table_prefix", "Table")
    
    # Header check
    headers = ["Field Name", "Utility", "Information", "Sample"]
    num_cols = len(headers)

    table = doc.add_table(rows=1 + len(field_details), cols=num_cols)
    table.autofit = False

    # Apply style borders
    border_color = style.tables.get("border_color", "CCCCCC")
    set_table_borders(table, border_color)

    # Set column widths
    table.columns[0].width = Inches(1.8)  # Field Name
    table.columns[1].width = Inches(2.2)  # Utility
    table.columns[2].width = Inches(1.5)  # Information
    table.columns[3].width = Inches(1.0)  # Sample

    # Render Header
    hdr_bg = style.tables.get("header_bg", style.table_header_bg)
    hdr_fg = style.tables.get("header_fg", style.table_header_fg)
    hdr_bold = style.tables.get("header_bold", True)

    hdr_cells = table.rows[0].cells
    for idx, col_name in enumerate(headers):
        hdr_cells[idx].text = col_name
        set_cell_background(hdr_cells[idx], hdr_bg)
        set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.runs:
            r = p.runs[0]
            r.font.name = style.heading_font
            r.font.size = Pt(9.5)
            r.font.bold = hdr_bold
            r.font.color.rgb = hex_to_rgb(hdr_fg)

    # Render Rows
    for r_idx, field in enumerate(field_details):
        row_cells = table.rows[r_idx + 1].cells
        
        # Split prefix out of the name if present (e.g. "Form Panel -> Field" -> "Field")
        raw_name = field.get("field_name", "Unnamed Field")
        display_name = raw_name.split(" -> ")[-1] if " -> " in raw_name else raw_name

        row_cells[0].text = display_name
        row_cells[1].text = field.get("utility", "")
        row_cells[2].text = field.get("information", "")
        row_cells[3].text = field.get("sample", "")

        # Style cells
        zebra_striping = style.tables.get("zebra_striping", True)
        zebra_color = style.get_color(style.tables.get("zebra_color", "table_zebra"))
        fill_color = zebra_color if (zebra_striping and r_idx % 2 == 1) else "FFFFFF"

        for idx, cell in enumerate(row_cells):
            set_cell_background(cell, fill_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            if p.runs:
                r = p.runs[0]
                r.font.name = style.body_font
                r.font.size = Pt(9)
                r.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

    # W4 Table Caption using native SEQ fields
    caption_text = field_details[0].get("field_name", "").split(" -> ")[0] if field_details else "Field Details"
    if "Panel" not in caption_text and "Table" not in caption_text and "Form" not in caption_text:
        caption_text = f"{caption_text} Table"

    from docbot.export.word_fields import add_caption
    module_num = int(table_number_str.split("-")[0]) if "-" in table_number_str else None
    
    add_caption(
        doc,
        prefix=table_prefix,
        caption_text=caption_text,
        style_cfg=style,
        seq_name="Table",
        module_num=module_num if style.numbering.get("table_format", "{module}-{tbl}").startswith("{module}") else None,
        font_size_pt=style.tables.get("caption_size_pt", 10),
    )


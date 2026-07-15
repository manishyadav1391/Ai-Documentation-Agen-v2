"""Table of Contents and Index listings renderers.

OOXML correctness note (T5.6):
  fldChar and instrText are **run-level** elements — they MUST be children of
  a ``<w:r>``, never direct children of ``<w:p>``.  Appending them straight
  to the paragraph element produces a structurally invalid document that
  LibreOffice silently tolerates but Microsoft Word refuses to open.

  Each field must look like:
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText xml:space="preserve">TOC …</w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>placeholder text</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
"""

from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from manual_builder.utils import add_styled_heading, add_body_paragraph


# ---------------------------------------------------------------------------
# Internal helper — produces a correctly-wrapped Word field inside *para*
# ---------------------------------------------------------------------------

def _insert_field(para, instr_text: str, placeholder: str = "") -> None:
    """
    Append a Word complex field to *para*, with each fldChar/instrText
    correctly wrapped inside its own ``<w:r>``.

    Args:
        para:         A python-docx Paragraph object.
        instr_text:   The raw field instruction, e.g. ``'TOC \\o "1-3" \\h \\z \\u'``.
        placeholder:  Human-readable text shown before Word updates the field.
    """
    def _run_with(el):
        """Create a bare run, append *el* to it, then append the run to the paragraph."""
        r = OxmlElement("w:r")
        r.append(el)
        para._p.append(r)
        return r

    # begin
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    _run_with(begin)

    # instruction
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    _run_with(instr)

    # separate
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    _run_with(sep)

    # placeholder run (added via python-docx so it inherits paragraph style)
    if placeholder:
        para.add_run(placeholder)

    # end
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    _run_with(end)


# ---------------------------------------------------------------------------
# Public renderers
# ---------------------------------------------------------------------------

def render_toc(doc, section_entry, manifest, style):
    """
    Adds the Table of Contents section with a native Word TOC field.
    Every fldChar/instrText is wrapped in its own <w:r> (OOXML-correct).
    """
    title_text = style.toc.get("title", "Table of Contents")
    add_styled_heading(doc, title_text, level=1, style_config=style)

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(12)
    r_desc = p_desc.add_run(
        "This Table of Contents is dynamically generated using native Microsoft Word index fields. "
        "Please right-click the block below and select \u2018Update Field\u2019 to refresh page numbers."
    )
    r_desc.font.name = style.body_font
    r_desc.font.size = Pt(9.5)
    r_desc.font.italic = True
    r_desc.font.color.rgb = RGBColor(110, 110, 110)

    max_depth = style.toc.get("max_depth", 3)
    p_toc = doc.add_paragraph()
    _insert_field(
        p_toc,
        instr_text=f'TOC \\o "1-{max_depth}" \\h \\z \\u',
        placeholder="[Right-click here and select \u2018Update Field\u2019 to generate Table of Contents]",
    )

    doc.add_page_break()


def render_table_of_tables(doc, section_entry, manifest, style):
    """
    Renders Table of Tables section with a native Word TOC \\c field.
    Every fldChar/instrText is wrapped in its own <w:r> (OOXML-correct).
    """
    add_styled_heading(doc, section_entry.heading or "Table of Tables", level=1, style_config=style)

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(12)
    r_desc = p_desc.add_run(
        "Please right-click the block below and select \u2018Update Field\u2019 to refresh Table of Tables listing."
    )
    r_desc.font.name = style.body_font
    r_desc.font.size = Pt(9.5)
    r_desc.font.italic = True
    r_desc.font.color.rgb = RGBColor(110, 110, 110)

    table_prefix = style.numbering.get("table_prefix", "Table")
    p_tot = doc.add_paragraph()
    _insert_field(
        p_tot,
        instr_text=f'TOC \\c "{table_prefix}"',
        placeholder=f"[Right-click here and select \u2018Update Field\u2019 to generate listing of captioned {table_prefix}s]",
    )

    doc.add_page_break()


def render_table_of_figures(doc, section_entry, manifest, style):
    """
    Renders Table of Figures section with a native Word TOC \\c field.
    Every fldChar/instrText is wrapped in its own <w:r> (OOXML-correct).
    """
    add_styled_heading(doc, section_entry.heading or "Table of Figures", level=1, style_config=style)

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(12)
    r_desc = p_desc.add_run(
        "Please right-click the block below and select \u2018Update Field\u2019 to refresh Table of Figures listing."
    )
    r_desc.font.name = style.body_font
    r_desc.font.size = Pt(9.5)
    r_desc.font.italic = True
    r_desc.font.color.rgb = RGBColor(110, 110, 110)

    figure_prefix = style.numbering.get("figure_prefix", "Figure")
    p_tof = doc.add_paragraph()
    _insert_field(
        p_tof,
        instr_text=f'TOC \\c "{figure_prefix}"',
        placeholder=f"[Right-click here and select \u2018Update Field\u2019 to generate listing of captioned {figure_prefix}s]",
    )

    doc.add_page_break()

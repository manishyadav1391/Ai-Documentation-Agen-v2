"""
Screen section renderer.

Changes (Spec E, F):
- E2: keep_with_next on image paragraph, keep_lines_together on caption.
- E3: 6.2" width, scale by height for tall captures, 0.5pt gray border via single-cell table.
- E4: Removed empty spacer paragraphs between sections.
- F1: Bullet field format: bold field name, en-dash, sentence-case description.
      Lead-in: "Enter the following details:"
- F2: Steps as numbered list (List Number style), bold UI labels in step text.
- F4: Caption format: "Figure N: Screen Name — caption_note" (colon form).
- F5: Skip/flag LLM-generated placeholder text in fields and purpose.
"""

from pathlib import Path
from typing import Any, Dict, List

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage

from manual_builder.utils import add_styled_heading, add_body_paragraph, hex_to_rgb
from manual_builder.renderers.field_table import render_field_table
from docbot.export.word_fields import add_caption

# Spec F5: LLM placeholder strings that must never reach the deliverable
_PLACEHOLDER_STRINGS = {"sample text", "data input", "lorem ipsum", "lorem", "["}


def _is_placeholder(text: str) -> bool:
    """Return True if the text looks like unresolved placeholder content."""
    if not text:
        return True
    lower = text.strip().lower()
    return any(p in lower for p in _PLACEHOLDER_STRINGS)


def _add_image_with_border(doc, img_path: Path, style) -> None:
    """
    E3: Insert screenshot inside a single-cell borderless-except-outline table
    so it gets a thin 0.5pt gray border around it.
    Also sets keep_with_next on the wrapper paragraph.
    """
    figures_cfg = style.raw.get("figures", {})
    max_w = figures_cfg.get("max_width_inches", 6.2)
    max_h = figures_cfg.get("max_height_inches", 7.5)

    # Determine render dimensions preserving aspect ratio
    fit_w, fit_h = max_w, max_w  # fallback
    try:
        with PILImage.open(img_path) as img:
            orig_w, orig_h = img.size
        aspect = orig_w / orig_h
        if orig_h / orig_w > 1.5:
            # Tall full-page capture: scale by height
            fit_h = min(max_h, max_w / aspect)
            fit_w = fit_h * aspect
        else:
            # Normal screenshot: scale by width
            fit_w = min(max_w, max_h * aspect)
            fit_h = fit_w / aspect
        # Safety caps
        if fit_h > max_h:
            fit_h = max_h
            fit_w = fit_h * aspect
        if fit_w > max_w:
            fit_w = max_w
            fit_h = fit_w / aspect
    except Exception:
        fit_w, fit_h = max_w, max_w / 1.5

    border_enabled = figures_cfg.get("border_enabled", True)
    border_color = figures_cfg.get("border_color", "AAAAAA").lstrip("#")
    border_pt = figures_cfg.get("border_pt", 0.5)
    border_sz = max(1, int(border_pt * 8))  # EMU: 1/8 pt per unit

    if border_enabled:
        # Wrap in a 1×1 table that provides the thin border
        from manual_builder.utils import set_table_borders, set_cell_margins, remove_table_borders
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        border_table = doc.add_table(rows=1, cols=1)
        border_table.autofit = False
        border_table.columns[0].width = Inches(fit_w + 0.05)
        remove_table_borders(border_table)

        # Apply thin border to the single cell
        cell = border_table.cell(0, 0)
        set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
        tcPr = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement as _El
        tcBorders = _El("w:tcBorders")
        for side in ["top", "left", "bottom", "right"]:
            b = _El(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(border_sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), border_color)
            tcBorders.append(b)
        tcPr.append(tcBorders)

        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(4)
        p_img.paragraph_format.keep_with_next = True  # E2

        try:
            p_img.add_run().add_picture(str(img_path), width=Inches(fit_w), height=Inches(fit_h))
        except Exception:
            p_img.add_run().add_picture(str(img_path), width=Inches(fit_w))
    else:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(4)
        p_img.paragraph_format.keep_with_next = True  # E2
        try:
            p_img.add_run().add_picture(str(img_path), width=Inches(fit_w), height=Inches(fit_h))
        except Exception:
            p_img.add_run().add_picture(str(img_path), width=Inches(fit_w))


def render_screen(doc, screen_index, session_dir, content_data, screen_meta, style, numbering):
    """
    Renders a screen section including headings, purpose, breadcrumbs,
    screenshots, captions, field list (bullets), steps, notes.
    """
    module_num = numbering.current_module

    # 1. Screen Title / Name
    screen_name = (
        screen_meta.get("screen_name") or
        content_data.get("screen_name") or
        screen_meta.get("h1_text") or
        screen_meta.get("title") or
        f"Screen {screen_index}"
    )

    screen_number_str = numbering.enter_section(level=2)
    add_styled_heading(doc, screen_name, level=2, style_config=style, numbering=screen_number_str)

    # 2. Screen Purpose (F5: skip placeholder)
    purpose = content_data.get("purpose")
    if not purpose and "screen_documentation" in content_data:
        purpose = content_data["screen_documentation"].get("overview")

    if purpose and not _is_placeholder(purpose):
        add_body_paragraph(
            doc, purpose,
            font_name=style.body_font, size_pt=style.body_size,
            color_hex=style.get_color("body_text"),
        )

    # 3. Path Breadcrumb
    path = content_data.get("path") or screen_meta.get("breadcrumb")
    if path:
        p_path = doc.add_paragraph()
        p_path.paragraph_format.space_before = Pt(2)
        p_path.paragraph_format.space_after = Pt(4)
        r_lbl = p_path.add_run("Path: ")
        r_lbl.font.name = style.body_font
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(style.body_size - 1)
        r_lbl.font.color.rgb = hex_to_rgb(style.get_color("secondary"))
        r_val = p_path.add_run(path)
        r_val.font.name = style.body_font
        r_val.font.size = Pt(style.body_size - 1)
        r_val.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

    # 4. Navigation instructions
    nav_instructions = content_data.get("navigation_instructions")
    if not nav_instructions and path:
        nav_instructions = (
            f"Select {screen_name} option from "
            f"{path.split(' >> ')[0] if ' >> ' in path else 'menu'} as shown in the image below."
        )
    if nav_instructions and not _is_placeholder(nav_instructions):
        add_body_paragraph(
            doc, nav_instructions,
            font_name=style.body_font, size_pt=style.body_size,
            color_hex=style.get_color("body_text"),
        )

    # 5. Screen Figures (Screenshots)
    figures_list = content_data.get("figures", [])
    if not figures_list:
        img_name = f"screen_{screen_index}_annotated.png"
        img_path = session_dir / img_name
        if not img_path.exists():
            img_path = session_dir / f"screen_{screen_index}.png"
        if img_path.exists():
            figures_list = [{"index": 1, "path": img_path.name, "caption_note": ""}]

    figure_prefix = style.numbering.get("figure_prefix", "Figure")

    for fig_entry in figures_list:
        fig_rel_path = fig_entry.get("path", "")
        # Prefer annotated image when it exists
        if fig_entry.get("index") == 1 or (fig_rel_path and fig_rel_path.startswith("screen_")):
            annotated_name = f"screen_{screen_index}_annotated.png"
            if (session_dir / annotated_name).exists():
                fig_rel_path = annotated_name

        fig_path = session_dir / fig_rel_path
        if not fig_path.exists():
            continue

        # E2/E3: image with optional border; keep_with_next set inside helper
        _add_image_with_border(doc, fig_path, style)

        # F4: caption format "Figure N: Screen Name — caption_note"
        caption_note = fig_entry.get("caption_note", "")
        if caption_note:
            fig_caption = f"{screen_name} \u2014 {caption_note}"
        else:
            fig_caption = screen_name

        fig_number_str = numbering.next_figure(module_num)

        # E2: Caption paragraph gets keep_lines_together
        # add_caption handles paragraph creation — we patch it after
        add_caption(
            doc,
            prefix=figure_prefix,
            caption_text=fig_caption,
            style_cfg=style,
            seq_name="Figure",
            module_num=module_num if numbering.mode == "module_prefixed" else None,
            font_size_pt=style.raw.get("figures", {}).get("caption_size_pt", 9),
        )
        # Apply keep_lines_together to the just-added caption paragraph (E2)
        try:
            doc.paragraphs[-1].paragraph_format.keep_lines_together = True
        except Exception:
            pass

        numbering.register_figure(fig_number_str, fig_caption)

    # 6. Field Details (Bullets — spec F1)
    field_details = content_data.get("field_details", [])
    if not field_details and "field_descriptions" in content_data:
        old_fields = content_data.get("field_descriptions", [])
        field_details = [
            {
                "field_name": f.get("field_name", ""),
                "utility": f.get("description", ""),
                "information": "",
                "sample": "",
            }
            for f in old_fields
        ]

    field_style = style.raw.get("fields", {}).get("style", "table")

    if field_details:
        if field_style == "bullets":
            # F1: Lead-in sentence
            lead_in = style.raw.get("fields", {}).get("lead_in", "Enter the following details:")
            p_hdr = doc.add_paragraph()
            p_hdr.paragraph_format.space_before = Pt(6)
            p_hdr.paragraph_format.space_after = Pt(2)
            r_hdr = p_hdr.add_run(lead_in)
            r_hdr.font.name = style.body_font
            r_hdr.font.bold = True
            r_hdr.font.size = Pt(style.body_size)
            r_hdr.font.color.rgb = hex_to_rgb(style.get_color("secondary"))

            for f in field_details:
                raw_name = f.get("field_name", "")
                display_name = raw_name.split(" -> ")[-1] if " -> " in raw_name else raw_name
                utility = f.get("utility", "")

                # F5: skip placeholder utilities
                if not display_name or _is_placeholder(utility):
                    continue

                p_bullet = doc.add_paragraph(style="List Bullet")
                p_bullet.paragraph_format.space_after = Pt(2)

                # F1: bold field name
                r_name = p_bullet.add_run(display_name)
                r_name.font.name = style.body_font
                r_name.font.bold = True
                r_name.font.size = Pt(style.body_size)
                r_name.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

                # en-dash separator
                r_sep = p_bullet.add_run(" \u2013 ")
                r_sep.font.name = style.body_font
                r_sep.font.size = Pt(style.body_size)
                r_sep.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

                # description (sentence case, ending with period if missing)
                desc = utility.strip()
                if desc and not desc.endswith("."):
                    desc += "."
                r_desc = p_bullet.add_run(desc)
                r_desc.font.name = style.body_font
                r_desc.font.size = Pt(style.body_size)
                r_desc.font.color.rgb = hex_to_rgb(style.get_color("body_text"))
        else:
            tbl_number_str = numbering.next_table(module_num)
            render_field_table(doc, screen_index, field_details, style, tbl_number_str)
            caption_text = (
                field_details[0].get("field_name", "").split(" -> ")[0]
                if field_details else "Field Details"
            )
            numbering.register_table(tbl_number_str, caption_text)

    # 7. Steps (F2: numbered list, bold UI labels)
    steps_list = content_data.get("steps", [])
    if steps_list:
        p_act_hdr = doc.add_paragraph()
        p_act_hdr.paragraph_format.space_before = Pt(6)
        p_act_hdr.paragraph_format.space_after = Pt(2)
        r_act_hdr = p_act_hdr.add_run("Steps:")
        r_act_hdr.font.name = style.body_font
        r_act_hdr.font.bold = True
        r_act_hdr.font.size = Pt(style.body_size)
        r_act_hdr.font.color.rgb = hex_to_rgb(style.get_color("secondary"))

        for step in steps_list:
            text = step.get("text", "")
            crop_path_rel = step.get("crop_path")
            crop_file = session_dir / crop_path_rel if crop_path_rel else None

            # F2: Use numbered list
            p_act = doc.add_paragraph(style="List Number")
            p_act.paragraph_format.space_after = Pt(2)

            # F2: Bold **text between asterisks** or <bold> markers
            _add_step_text_with_bold(p_act, text, style)

            if crop_file and crop_file.exists():
                try:
                    p_act.add_run(" ").add_picture(str(crop_file), height=Inches(0.25))
                except Exception:
                    pass
    else:
        # Fallback button/filter text
        screen_doc = content_data.get("screen_documentation", {}) or {}
        actions = []
        for btn in screen_doc.get("buttons", []):
            actions.append(f"Click the **{btn}** button to perform the corresponding action.")
        for sf in screen_doc.get("search_filters", []):
            actions.append(f"Filter records by entering details in the **{sf}** field.")

        if actions:
            p_act_hdr = doc.add_paragraph()
            p_act_hdr.paragraph_format.space_before = Pt(6)
            p_act_hdr.paragraph_format.space_after = Pt(2)
            r_act_hdr = p_act_hdr.add_run("Steps:")
            r_act_hdr.font.name = style.body_font
            r_act_hdr.font.bold = True
            r_act_hdr.font.size = Pt(style.body_size)
            r_act_hdr.font.color.rgb = hex_to_rgb(style.get_color("secondary"))

            for act in actions[:6]:
                p_act = doc.add_paragraph(style="List Number")
                p_act.paragraph_format.space_after = Pt(2)
                _add_step_text_with_bold(p_act, act, style)

    # 8. Notes list (F3)
    notes = content_data.get("notes") or (content_data.get("screen_documentation", {}) or {}).get("notes", [])
    if notes:
        for note in notes:
            if not note or _is_placeholder(note):
                continue
            p_note = doc.add_paragraph()
            p_note.paragraph_format.space_before = Pt(4)
            p_note.paragraph_format.space_after = Pt(4)
            r_note_lbl = p_note.add_run("Note: ")
            r_note_lbl.font.name = style.body_font
            r_note_lbl.font.bold = True
            r_note_lbl.font.size = Pt(style.body_size)
            r_note_lbl.font.color.rgb = hex_to_rgb(style.get_color("tertiary"))
            r_note_val = p_note.add_run(note)
            r_note_val.font.name = style.body_font
            r_note_val.font.size = Pt(style.body_size)
            r_note_val.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

    # E1: Configurable page break per screen
    page_break = style.raw.get("layout", {}).get("page_break_per_screen", False)
    if page_break:
        doc.add_page_break()


def _add_step_text_with_bold(paragraph, text: str, style) -> None:
    """
    F2: Render step text, bolding any token wrapped in **…** (markdown-like bold).
    Adds runs with alternating bold/normal formatting.
    """
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(part)
            r.bold = False
        r.font.name = style.body_font
        r.font.size = Pt(style.body_size)
        r.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

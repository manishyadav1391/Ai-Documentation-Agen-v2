"""
Cover page renderer.

Design (per current spec):
- No page header on the cover (relies on section.different_first_page_header_footer,
  which is already set to True in utils.setup_header_footer — so page 1 gets an
  empty first_page_header automatically).
- Banner block embedded as the first body element. Supports either:
    (a) two images side-by-side (left_image_path + right_image_path in style.cover), or
    (b) a single wide banner (banner_path in style.cover).
- Title lines: fully data-driven list from cover.yaml → title_lines: [...].
  Each line rendered centered, bold, primary color. Falls back to sensible
  defaults from the manifest if title_lines is not provided.
- Version-only footer block near the bottom (no audience, no metadata table).
- Everything sized/spaced from style.cover so different clients can retune
  without code changes.

Spec preserved:
- B2: fail-loud if client_display_name is missing / placeholder.
- D3: skip entirely when manifest.cover_enabled is False.
"""

from pathlib import Path

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from manual_builder.utils import (
    remove_table_borders,
    hex_to_rgb,
)
from manual_builder.build_error import BuildError


def _resolve_image(rel_path: str, manifest) -> Path | None:
    """Resolve an image path relative to the client's content dir, then absolute."""
    if not rel_path:
        return None
    p = manifest.get_source_path(rel_path)
    if p.exists():
        return p
    p2 = Path(rel_path)
    return p2 if p2.exists() else None


def render_cover(doc, section_entry, manifest, style):
    """
    Render a clean, data-driven cover page.

    Style keys (all under ``style.cover`` in <client>/style.yaml):
        banner_path              — single banner image (relative to content_dir)
        left_image_path          — left half of a 2-up banner
        right_image_path         — right half of a 2-up banner
        banner_height_inches     — height for banner images (default 0.75)
        title_color              — color name (default "primary") or hex
        title_size_pt            — main title lines (default 26)
        title_line_spacing       — line spacing for the title block (default 1.5)
        title_space_after_pt     — spacing after each title line (default 8)
        version_label            — label before the version number (default "Version")
        version_size_pt          — version line font size (default 18)
        version_color            — color for version line (default = title_color)
        top_spacer_paragraphs    — blank paragraphs between banner and title (default 3)
        bottom_spacer_paragraphs — blank paragraphs between title and version (default 6)

    Cover-content keys (under the file pointed to by section.source, typically
    ``clients/<key>/content/cover.yaml``):
        title_lines: [ "...", "...", ... ]     — one entry per centered title line
        version_label: "Version"               — overrides style version_label
    """

    # D3: honour cover_enabled flag
    if not manifest.cover_enabled:
        return

    # B2: validate required cover fields
    client_display_name = manifest.client_display_name
    if not client_display_name or client_display_name.strip().startswith("["):
        raise BuildError(
            "manifest field 'client_display_name' is required by the cover renderer"
        )

    # ── Load optional cover.yaml override data ────────────────────────────
    cover_data = {}
    if section_entry.source:
        cover_path = manifest.get_source_path(section_entry.source)
        if cover_path.exists():
            import yaml
            with cover_path.open("r", encoding="utf-8") as f:
                cover_data = yaml.safe_load(f) or {}

    cover_style = style.cover

    # ── Resolve title lines (data-driven with sensible fallbacks) ─────────
    title_lines = cover_data.get("title_lines") or []
    if not title_lines:
        # Fallback: build a sane default from manifest so clients that haven't
        # migrated to cover.yaml still get a usable cover.
        fallback_document_title = cover_data.get("document_title", "User Manual")
        title_lines = [
            manifest.client_display_name,
            manifest.system_name or "",
            fallback_document_title,
        ]
    # Drop empty lines but preserve order
    title_lines = [str(line).strip() for line in title_lines if line and str(line).strip()]

    # ── Style knobs ───────────────────────────────────────────────────────
    title_color_hex = style.get_color(cover_style.get("title_color", "primary"))
    version_color_hex = style.get_color(cover_style.get("version_color",
                                                        cover_style.get("title_color", "primary")))
    title_size_pt = cover_style.get("title_size_pt", 26)
    title_line_spacing = cover_style.get("title_line_spacing", 1.5)
    title_space_after_pt = cover_style.get("title_space_after_pt", 8)

    version_label = cover_data.get("version_label",
                                   cover_style.get("version_label", "Version"))
    version_value = manifest.document_version or manifest.version or "1.0"
    version_size_pt = cover_style.get("version_size_pt", 18)

    top_spacer = int(cover_style.get("top_spacer_paragraphs", 3))
    bottom_spacer = int(cover_style.get("bottom_spacer_paragraphs", 6))
    banner_height_in = float(cover_style.get("banner_height_inches", 0.75))

    # ── Banner row (as first body element) ────────────────────────────────
    left_img = _resolve_image(cover_style.get("left_image_path", ""), manifest)
    right_img = _resolve_image(cover_style.get("right_image_path", ""), manifest)
    single_banner = _resolve_image(cover_style.get("banner_path", ""), manifest)

    if left_img or right_img:
        # Two-image side-by-side layout (matches ministry-logo + program-banner style)
        banner_table = doc.add_table(rows=1, cols=2)
        banner_table.autofit = False
        banner_table.columns[0].width = Inches(3.25)
        banner_table.columns[1].width = Inches(3.25)
        remove_table_borders(banner_table)

        left_cell = banner_table.cell(0, 0)
        p_l = left_cell.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_l.paragraph_format.space_after = Pt(0)
        if left_img:
            try:
                p_l.add_run().add_picture(str(left_img), height=Inches(banner_height_in))
            except Exception:
                pass

        right_cell = banner_table.cell(0, 1)
        p_r = right_cell.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_after = Pt(0)
        if right_img:
            try:
                p_r.add_run().add_picture(str(right_img), height=Inches(banner_height_in))
            except Exception:
                pass
    elif single_banner:
        # Single wide banner
        p_banner = doc.add_paragraph()
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_banner.paragraph_format.space_after = Pt(0)
        try:
            p_banner.add_run().add_picture(str(single_banner), width=Inches(6.5))
        except Exception:
            pass

    # ── Spacer between banner and title ──────────────────────────────────
    for _ in range(top_spacer):
        doc.add_paragraph()

    # ── Title lines (all centered, primary color, bold) ──────────────────
    for line in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(title_space_after_pt)
        p.paragraph_format.line_spacing = title_line_spacing
        r = p.add_run(line)
        r.font.name = style.heading_font
        r.font.size = Pt(title_size_pt)
        r.font.bold = True
        r.font.color.rgb = hex_to_rgb(title_color_hex)

    # ── Push version block toward bottom ─────────────────────────────────
    for _ in range(bottom_spacer):
        doc.add_paragraph()

    # ── Version line (only) ──────────────────────────────────────────────
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ver = p_ver.add_run(f"{version_label} {version_value}")
    r_ver.font.name = style.heading_font
    r_ver.font.size = Pt(version_size_pt)
    r_ver.font.color.rgb = hex_to_rgb(version_color_hex)

    doc.add_page_break()
"""Prose (markdown) renderer."""

import re
from docx.shared import Pt
from manual_builder.utils import add_styled_heading, add_body_paragraph, hex_to_rgb


def _apply_inline_formatting(run_target, text, style):
    """Processes markdown bold/italic tags and adds formatted runs to a paragraph."""
    # Split text by bold (** or __) and italic (* or _) markers
    parts = re.split(r'(\*\*|__|\*|_)', text)
    
    is_bold = False
    is_italic = False
    
    for part in parts:
        if part in ('**', '__'):
            is_bold = not is_bold
            continue
        elif part in ('*', '_'):
            is_italic = not is_italic
            continue
        
        if not part:
            continue
            
        r = run_target.add_run(part)
        r.font.name = style.body_font
        r.font.size = Pt(style.body_size)
        r.font.bold = is_bold
        r.font.italic = is_italic
        r.font.color.rgb = hex_to_rgb(style.get_color("body_text"))


def render_prose(doc, section_entry, manifest, style, heading_text: str = None):
    """
    Reads a markdown file, parses headings, bullets, bold/italic, and renders to docx.
    """
    # 1. Heading
    title = heading_text if heading_text is not None else section_entry.heading
    if title:
        add_styled_heading(doc, title, level=1, style_config=style)


    prose_path = manifest.get_source_path(section_entry.source)
    if not prose_path.exists():
        return

    content = prose_path.read_text(encoding="utf-8")
    
    # Process line by line
    lines = content.splitlines()
    in_list = False
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # Check for headings
        h_match = re.match(r'^(#{2,4})\s+(.*)$', stripped)
        if h_match:
            hashes, h_text = h_match.groups()
            level = len(hashes)  # ## is 2, ### is 3, #### is 4
            add_styled_heading(doc, h_text, level=level, style_config=style)
            in_list = False
            continue
            
        # Check for bullets
        b_match = re.match(r'^[\-\*]\s+(.*)$', stripped)
        if b_match:
            bullet_text = b_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            _apply_inline_formatting(p, bullet_text, style)
            in_list = True
            continue
            
        # Standard paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        _apply_inline_formatting(p, stripped, style)
        in_list = False

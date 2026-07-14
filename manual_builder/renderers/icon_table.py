"""Icon and Button reference table renderer."""

from pathlib import Path
from typing import Any, Dict
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import yaml
from manual_builder.utils import add_styled_heading, set_cell_background, set_cell_margins, set_table_borders, hex_to_rgb


def render_icon_table(doc, section_entry, manifest, style):
    """
    Renders reference tables for buttons/icons.
    If icons/images are skipped or not found, renders a clean text-based table instead.
    """
    if section_entry.heading:
        add_styled_heading(doc, section_entry.heading, level=2, style_config=style)

    table_path = manifest.get_source_path(section_entry.source)
    if not table_path.exists():
        return

    with table_path.open("r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}

    columns = data.get("columns", ["Name", "Description"])
    rows = data.get("rows", [])
    if not rows:
        return

    num_cols = len(columns)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.autofit = False

    # Apply style borders
    border_color = style.tables.get("border_color", "CCCCCC")
    set_table_borders(table, border_color)

    # Set column widths based on column count
    if num_cols == 3:
        table.columns[0].width = Inches(1.5)  # Button/Icon
        table.columns[1].width = Inches(1.8)  # Name
        table.columns[2].width = Inches(3.2)  # Description
    else:
        table.columns[0].width = Inches(2.2)
        table.columns[1].width = Inches(4.3)

    # Render Header
    hdr_bg = style.tables.get("header_bg", style.table_header_bg)
    hdr_fg = style.tables.get("header_fg", style.table_header_fg)
    hdr_bold = style.tables.get("header_bold", True)

    hdr_cells = table.rows[0].cells
    for idx, col_name in enumerate(columns):
        hdr_cells[idx].text = col_name
        set_cell_background(hdr_cells[idx], hdr_bg)
        set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.runs:
            r = p.runs[0]
            r.font.name = style.heading_font
            r.font.size = Pt(9.5)
            r.font.bold = hdr_bold
            r.font.color.rgb = hex_to_rgb(hdr_fg)

    # Render Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        
        # Determine cell values based on columns structure
        cell_values = []
        if num_cols == 3:
            # First column is placeholder for the button/icon
            cell_values.append("") 
            cell_values.append(str(row_data.get("name", "")))
            cell_values.append(str(row_data.get("description", "")))
        else:
            cell_values.append(str(row_data.get("name", "")))
            cell_values.append(str(row_data.get("description", "")))

        # Fill text
        for idx, val in enumerate(cell_values):
            row_cells[idx].text = val

        # Handle button/icon picture if path exists (otherwise remains blank text)
        if num_cols == 3 and "icon" in row_data:
            icon_path_str = row_data["icon"]
            if icon_path_str:
                icon_path = Path(icon_path_str)
                if not icon_path.is_absolute() and manifest:
                    icon_path = manifest.get_source_path(icon_path_str)
                if icon_path.exists():
                    p_img = row_cells[0].paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p_img.add_run().add_picture(str(icon_path), height=Inches(0.25))
                    except Exception:
                        # Fallback to name in first col if img fails
                        row_cells[0].text = f"[{row_data.get('name', 'Icon')}]"

        # Apply cell shading and fonts
        zebra_striping = style.tables.get("zebra_striping", True)
        zebra_color = style.get_color(style.tables.get("zebra_color", "table_zebra"))
        fill_color = zebra_color if (zebra_striping and r_idx % 2 == 1) else "FFFFFF"

        for idx, cell in enumerate(row_cells):
            set_cell_background(cell, fill_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            if p.runs:
                r = p.runs[0]
                r.font.name = style.body_font
                r.font.size = Pt(9)
                r.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

    # Spacer after table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(8)

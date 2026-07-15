"""
Revision history renderer.

Spec B4: Renders the revision history table.
- Columns: Version | Date | Description | Prepared By | Reviewed By | Approved By
- If the entries list is empty, auto-inserts one seed row using manifest metadata.
- Date field: if blank, fills with today's date (ISO format).
"""

import datetime
from pathlib import Path
from typing import Any, Dict

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from manual_builder.utils import (
    set_cell_background, set_cell_margins, set_table_borders, hex_to_rgb, add_styled_heading
)


def render_revision_history(doc, section_entry, manifest, style):
    """
    Renders the revision history table.

    Column set matches reference manual:
    Version | Date | Description | Prepared By | Reviewed By | Approved By
    """
    add_styled_heading(doc, section_entry.heading or "Revision History", level=1, style_config=style)

    # Load data from source file
    history_path = manifest.get_source_path(section_entry.source) if section_entry.source else None
    import yaml

    data = {}
    if history_path and history_path.exists():
        with history_path.open("r", encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}

    entries = data.get("entries", [])

    # B4: auto-seed one row if entries list is empty
    if not entries:
        entries = [{
            "version": manifest.document_version or manifest.version or "1.0",
            "date": "",
            "description": "Initial version",
            "prepared_by": manifest.prepared_by or "",
            "reviewed_by": manifest.reviewed_by or "",
            "approved_by": manifest.approved_by or "",
        }]

    # Columns from style, defaulting to the reference-manual set
    headers = style.revision_history.get("columns") or [
        "Version", "Date", "Description", "Prepared By", "Reviewed By", "Approved By"
    ]

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(entries), cols=num_cols)
    table.autofit = False

    border_color = style.revision_history.get("border_color", "CCCCCC")
    set_table_borders(table, border_color)

    # Evenly distribute width across 6.2" usable content width
    col_width = Inches(6.2 / num_cols)
    for c_idx in range(num_cols):
        table.columns[c_idx].width = col_width

    hdr_bg = style.revision_history.get("header_bg", style.table_header_bg)
    hdr_fg = style.revision_history.get("header_fg", style.table_header_fg)

    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = hdr_cells[idx]
        cell.text = text
        set_cell_background(cell, hdr_bg)
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.runs:
            r = p.runs[0]
            r.font.name = style.heading_font
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = hex_to_rgb(hdr_fg)

    today_str = datetime.date.today().strftime("%d-%m-%Y")

    def get_val(entry: dict, col_name: str) -> str:
        """Map column header name to entry dict key."""
        c = col_name.lower()
        if "version" in c:
            val = entry.get("version") or entry.get("revision") or "1.0"
        elif "date" in c:
            val = entry.get("date") or ""
            if not val:
                val = today_str  # auto-fill today's date
        elif "description" in c or "change" in c:
            val = entry.get("description") or entry.get("description_of_change") or ""
        elif "prepared" in c:
            val = entry.get("prepared_by") or entry.get("by") or entry.get("author") or ""
        elif "reviewed" in c:
            val = entry.get("reviewed_by") or ""
        elif "approved" in c:
            val = entry.get("approved_by") or ""
        else:
            val = entry.get(col_name) or entry.get(col_name.lower().replace(" ", "_")) or ""
        return str(val)

    for r_idx, entry in enumerate(entries):
        row_cells = table.rows[r_idx + 1].cells
        fill_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"

        for idx, col_name in enumerate(headers):
            cell = row_cells[idx]
            cell.text = get_val(entry, col_name)
            set_cell_background(cell, fill_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            if p.runs:
                r = p.runs[0]
                r.font.name = style.body_font
                r.font.size = Pt(9)
                r.font.color.rgb = hex_to_rgb(style.get_color("body_text"))

    doc.add_page_break()

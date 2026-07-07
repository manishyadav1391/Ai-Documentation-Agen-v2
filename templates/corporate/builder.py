from pathlib import Path
from typing import Any, Dict
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from templates.base import DocumentBuilder

def hex_to_rgb(hex_str: str) -> RGBColor:
    """Helper to convert a hex color string to docx RGBColor."""
    hex_str = hex_str.lstrip('#')
    if len(hex_str) != 6:
        # Default back to a dark slate
        return RGBColor(27, 54, 93)
    return RGBColor(*(int(hex_str[i:i+2], 16) for i in (0, 2, 4)))

class CorporateBuilder(DocumentBuilder):
    """
    Assembles a highly professional, corporate-styled manual.
    All fonts, colors, and layout elements adapt dynamically to user preferences.
    """
    def __init__(self, style_config: Dict[str, Any] = None):
        super().__init__(style_config)

    def _setup_styles(self):
        """Initializes custom font family, heading styles, margins, and document-wide header/footer."""
        self.set_margins(1.0, 1.1, 1.0, 1.0)  # slightly more bottom margin for footer

        # Read styling configurations
        self.font_name = self.style_config.get("font_name", "Segoe UI")
        self.primary_hex = self.style_config.get("primary_color", "1B365D").lstrip('#')
        self.secondary_hex = self.style_config.get("secondary_color", "D97706").lstrip('#')
        self.text_hex = self.style_config.get("text_color", "333333").lstrip('#')

        # Normal style
        normal = self.doc.styles['Normal']
        normal.font.name = self.font_name
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = hex_to_rgb(self.text_hex)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        # Heading styles
        heading_specs = {
            'Heading 1': (Pt(16), True, Pt(18), Pt(6)),
            'Heading 2': (Pt(13), True, Pt(14), Pt(4)),
            'Heading 3': (Pt(11), True, Pt(10), Pt(3)),
        }
        for h_name, (size, bold, space_before, space_after) in heading_specs.items():
            if h_name in self.doc.styles:
                h_style = self.doc.styles[h_name]
                h_style.font.name = self.font_name
                h_style.font.size = size
                h_style.font.bold = bold
                h_style.font.color.rgb = hex_to_rgb(self.primary_hex)
                h_style.paragraph_format.space_before = space_before
                h_style.paragraph_format.space_after = space_after
                h_style.paragraph_format.keep_with_next = True

        # ── Set document-wide header/footer ONCE here ─────────────────────────
        logo_val = self.style_config.get("logo_path", "")
        logo_path = Path(logo_val) if logo_val else None
        doc_title = self.style_config.get("company_name", "System Documentation") + " User Manual"
        self.setup_header_footer(logo_path, doc_title, self.font_name)

    def add_cover_page(self, title: str = "User Manual"):
        """Generates a premium corporate cover page. Cover page has no header/footer."""
        import datetime
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        logo_val = self.style_config.get("logo_path", "")
        logo_path = Path(logo_val) if logo_val else None
        company_name = self.style_config.get("company_name", "Corporate Corporation")
        subtitle_text = self.style_config.get("subtitle", "Comprehensive User Guide & Technical Documentation")
        app_name = self.style_config.get("app_name", "")

        # ── Decorative top accent bar (colored rectangle using a table) ────────
        accent_table = self.doc.add_table(rows=1, cols=1)
        accent_table.autofit = False
        accent_table.columns[0].width = Inches(6.5)
        accent_cell = accent_table.cell(0, 0)
        self.set_cell_background(accent_cell, self.primary_hex)
        self.set_cell_margins(accent_cell, top=200, bottom=200, left=150, right=150)
        accent_p = accent_cell.paragraphs[0]
        accent_p.paragraph_format.space_after = Pt(0)

        # ── Spacing ───────────────────────────────────────────────────────────
        for _ in range(3):
            self.doc.add_paragraph()

        # ── Logo (large, left-aligned) ────────────────────────────────────────
        if logo_path and logo_path.exists():
            p_logo = self.doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p_logo.add_run().add_picture(str(logo_path), height=Inches(1.1))
            except Exception:
                pass

        self.doc.add_paragraph()  # spacer

        # ── Company name (small caps, secondary accent) ────────────────────────
        p_comp = self.doc.add_paragraph()
        r_comp = p_comp.add_run(company_name.upper())
        r_comp.font.name = self.font_name
        r_comp.font.size = Pt(11)
        r_comp.font.bold = True
        r_comp.font.color.rgb = hex_to_rgb(self.secondary_hex)

        # ── Main title ────────────────────────────────────────────────────────
        p_title = self.doc.add_paragraph()
        r_title = p_title.add_run(title)
        r_title.font.name = self.font_name
        r_title.font.size = Pt(32)
        r_title.font.bold = True
        r_title.font.color.rgb = hex_to_rgb(self.primary_hex)
        self.add_bottom_border_to_paragraph(p_title, color_hex=self.secondary_hex, sz=18)

        # ── App name (optional, under title) ──────────────────────────────────
        if app_name:
            p_app = self.doc.add_paragraph()
            r_app = p_app.add_run(app_name)
            r_app.font.name = self.font_name
            r_app.font.size = Pt(14)
            r_app.font.italic = True
            r_app.font.color.rgb = hex_to_rgb(self.secondary_hex)

        # ── Subtitle ──────────────────────────────────────────────────────────
        p_sub = self.doc.add_paragraph()
        r_sub = p_sub.add_run(subtitle_text)
        r_sub.font.name = self.font_name
        r_sub.font.size = Pt(12)
        r_sub.font.italic = True
        r_sub.font.color.rgb = RGBColor(130, 130, 130)

        # ── Push metadata block to bottom ─────────────────────────────────────
        for _ in range(10):
            self.doc.add_paragraph()

        # ── Bottom decorative line ────────────────────────────────────────────
        p_line = self.doc.add_paragraph()
        self.add_bottom_border_to_paragraph(p_line, color_hex=self.primary_hex, sz=12)
        p_line.paragraph_format.space_after = Pt(8)

        # ── Metadata block ────────────────────────────────────────────────────
        meta_table = self.doc.add_table(rows=1, cols=2)
        meta_table.autofit = False
        meta_table.columns[0].width = Inches(3.0)
        meta_table.columns[1].width = Inches(3.5)
        self.set_table_borders(meta_table, "E2E8F0")

        left_cell = meta_table.cell(0, 0)
        self.set_cell_background(left_cell, "F8FAFC")
        self.set_cell_margins(left_cell, top=120, bottom=120, left=150, right=150)
        p_left = left_cell.paragraphs[0]
        p_left.paragraph_format.line_spacing = 1.5
        for label, value in [
            ("DOCUMENT VERSION", "1.0  (Official Draft)"),
            ("PUBLISHED DATE", datetime.datetime.now().strftime("%B %d, %Y")),
            ("GENERATED BY", "Documentation Automation Engine"),
        ]:
            r_lbl = p_left.add_run(f"{label}:\n")
            r_lbl.font.name = self.font_name
            r_lbl.font.size = Pt(8)
            r_lbl.font.bold = True
            r_lbl.font.color.rgb = RGBColor(100, 100, 100)
            r_val = p_left.add_run(f"{value}\n\n")
            r_val.font.name = self.font_name
            r_val.font.size = Pt(9)
            r_val.font.color.rgb = RGBColor(40, 40, 40)

        right_cell = meta_table.cell(0, 1)
        self.set_cell_background(right_cell, self.primary_hex)
        self.set_cell_margins(right_cell, top=120, bottom=120, left=150, right=150)
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right.paragraph_format.space_before = Pt(20)
        r_conf = p_right.add_run("CONFIDENTIAL")
        r_conf.font.name = self.font_name
        r_conf.font.size = Pt(10)
        r_conf.font.bold = True
        r_conf.font.color.rgb = RGBColor(255, 255, 255)

        self.doc.add_page_break()

    def add_divider_page(self, title: str, description: str = ""):
        """Creates a professional section divider page for a module, starting a new section."""
        new_section = self.doc.add_section()
        new_section.different_first_page_header_footer = True
        
        # Keep same margins for the new section
        new_section.top_margin = Inches(1.0)
        new_section.bottom_margin = Inches(1.1)
        new_section.left_margin = Inches(1.0)
        new_section.right_margin = Inches(1.0)

        # Build some vertical space
        for _ in range(5):
            self.doc.add_paragraph()

        # Large Section Heading
        p_sec = self.doc.add_paragraph()
        p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sec = p_sec.add_run(title)
        run_sec.font.name = self.font_name
        run_sec.font.size = Pt(26)
        run_sec.font.bold = True
        run_sec.font.color.rgb = hex_to_rgb(self.primary_hex)

        # Horizontal Divider line
        p_line = self.doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_bottom_border_to_paragraph(p_line, color_hex=self.secondary_hex, sz=12)

        # Optional short description
        if description:
            p_desc = self.doc.add_paragraph()
            p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_desc = p_desc.add_run(description)
            run_desc.font.name = self.font_name
            run_desc.font.size = Pt(11)
            run_desc.font.italic = True
            run_desc.font.color.rgb = RGBColor(110, 110, 110)

        self.doc.add_page_break()

    def add_revision_history(self):
        """Creates the revision control table styled to corporate brand."""
        h = self.doc.add_heading(level=1)
        run = h.add_run("Revision History")
        run.font.name = self.font_name
        run.font.bold = True
        run.font.color.rgb = hex_to_rgb(self.primary_hex)

        table = self.doc.add_table(rows=2, cols=5)
        table.autofit = False
        self.set_table_borders(table, "E2E8F0")

        table.columns[0].width = Inches(0.8)
        table.columns[1].width = Inches(1.1)
        table.columns[2].width = Inches(1.4)
        table.columns[3].width = Inches(1.4)
        table.columns[4].width = Inches(1.8)

        headers = ['Ver.', 'Date', 'Author', 'Approved By', 'Details of Changes']
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            self.set_cell_background(hdr_cells[idx], self.primary_hex)
            self.set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=120, right=120)
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.runs[0]
            r.font.name = self.font_name
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

        import datetime
        current_date = datetime.datetime.now().strftime("%Y-%m-%d")

        row_cells = table.rows[1].cells
        data = ["1.0", current_date, "Automation Bot", "Quality Lead", "Initial generated manual — Version 1.0"]
        for idx, val in enumerate(data):
            row_cells[idx].text = val
            self.set_cell_background(row_cells[idx], "F8FAFC")
            self.set_cell_margins(row_cells[idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[idx].paragraphs[0]
            if p.runs:
                p.runs[0].font.name = self.font_name
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = hex_to_rgb(self.text_hex)

        # Space after table + page break
        p_space = self.doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(8)
        self.doc.add_page_break()

    def add_draft_title(self, title: str, description: str = ""):
        """Adds a professional title block at the top of a module draft document."""
        p_title = self.doc.add_paragraph()
        run_title = p_title.add_run(title)
        run_title.font.name = self.font_name
        run_title.font.size = Pt(22)
        run_title.font.bold = True
        run_title.font.color.rgb = hex_to_rgb(self.primary_hex)
        self.add_bottom_border_to_paragraph(p_title, color_hex=self.secondary_hex, sz=12)
        p_title.paragraph_format.space_after = Pt(8)

        if description:
            p_desc = self.doc.add_paragraph()
            run_desc = p_desc.add_run(description)
            run_desc.font.name = self.font_name
            run_desc.font.size = Pt(10)
            run_desc.font.italic = True
            run_desc.font.color.rgb = RGBColor(110, 110, 110)
            p_desc.paragraph_format.space_after = Pt(20)

    def add_toc_placeholder(self):
        """Adds a Table of Contents placeholder styled with native MS Word TOC fields."""
        h = self.doc.add_heading(level=1)
        run = h.add_run("Table of Contents")
        run.font.name = self.font_name
        run.font.color.rgb = hex_to_rgb(self.primary_hex)
        
        # Add descriptive text
        p_desc = self.doc.add_paragraph()
        run_desc = p_desc.add_run("This Table of Contents is dynamically generated using native Microsoft Word index fields. Please right-click the block below and select 'Update Field' to refresh page numbers.")
        run_desc.font.size = Pt(9.5)
        run_desc.font.italic = True
        run_desc.font.color.rgb = RGBColor(110, 110, 110)
        p_desc.paragraph_format.space_after = Pt(12)

        p_toc = self.doc.add_paragraph()
        self.add_toc_field(p_toc)
        
        self.doc.add_page_break()

    def add_callout(self, text: str, callout_type: str = "note"):
        """Adds a professional left-bordered callout box to emphasize text."""
        # Convert type to formatting colors
        types = {
            "warning": {"color": "DC2626", "bg": "FEF2F2", "title": "WARNING"},
            "important": {"color": "D97706", "bg": "FFFBEB", "title": "IMPORTANT"},
            "note": {"color": self.primary_hex, "bg": "F8FAFC", "title": "NOTE"}
        }
        cfg = types.get(callout_type.lower(), types["note"])
        
        # Create a single cell table
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.0)
        
        cell = table.cell(0, 0)
        self.set_cell_background(cell, cfg["bg"])
        self.set_cell_left_border(cell, cfg["color"], size_sz=36)
        self.set_cell_margins(cell, top=140, bottom=140, left=180, right=140)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        
        # Add Callout Title
        run_title = p.add_run(f"{cfg['title']}: ")
        run_title.font.name = self.font_name
        run_title.font.bold = True
        run_title.font.size = Pt(10)
        run_title.font.color.rgb = hex_to_rgb(cfg["color"])
        
        # Add Callout Content
        run_text = p.add_run(text)
        run_text.font.name = self.font_name
        run_text.font.italic = True
        run_text.font.size = Pt(9.5)
        run_text.font.color.rgb = hex_to_rgb(self.text_hex)
        
        # Spacer paragraph after the table
        p_space = self.doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(4)
        p_space.paragraph_format.space_after = Pt(4)

    def add_screen_section(self, screen_index: int, image_path: Path, content_data: Dict[str, Any], screen_meta: Dict[str, Any] = None):
        """Renders a structured screen section with professional headings, prose sections, and element table."""
        screen_meta = screen_meta or {}

        # Derive semantic screen title from meta or content_data
        screen_name = (
            screen_meta.get("screen_name") or
            content_data.get("screen_name") or
            screen_meta.get("h1_text") or
            screen_meta.get("title") or
            f"Screen {screen_index}"
        )

        # Heading 1 — Screen Name (styled via Heading 1 style, plus bottom border)
        h = self.doc.add_heading(level=1)
        run = h.add_run(f"{screen_index}. {screen_name}")
        run.font.name = self.font_name
        run.font.bold = True
        run.font.color.rgb = hex_to_rgb(self.primary_hex)
        self.add_bottom_border_to_paragraph(h, color_hex=self.secondary_hex, sz=6)

        # Screen image with figure caption
        if image_path.exists():
            p_space = self.doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(4)
            p_space.paragraph_format.space_after = Pt(0)

            p_img = self.doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(4)

            # Determine dynamic size using PIL
            from PIL import Image as PILImage
            try:
                with PILImage.open(image_path) as img:
                    img_w, img_h = img.size
                aspect = img_w / img_h
                # Fit within max width of 6.2 inches and max height of 4.2 inches
                max_w = 6.2
                max_h = 4.2
                if aspect > (max_w / max_h):
                    # Width-bound
                    fit_w = max_w
                    fit_h = max_w / aspect
                else:
                    # Height-bound
                    fit_h = max_h
                    fit_w = max_h * aspect
                p_img.add_run().add_picture(str(image_path), width=Inches(fit_w), height=Inches(fit_h))
            except Exception:
                # Fallback to standard width-only
                p_img.add_run().add_picture(str(image_path), width=Inches(5.8))

            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(4)
            cap.paragraph_format.space_after = Pt(14)
            
            r_fig = cap.add_run(f"Figure {screen_index}: ")
            r_fig.font.name = self.font_name
            r_fig.font.bold = True
            r_fig.font.size = Pt(9.5)
            r_fig.font.color.rgb = hex_to_rgb(self.secondary_hex)

            r_cap = cap.add_run(screen_name)
            r_cap.font.name = self.font_name
            r_cap.font.italic = True
            r_cap.font.size = Pt(9.5)
            r_cap.font.color.rgb = RGBColor(100, 100, 100)


        # Procedure Prose — parsed into structured sections
        prose = content_data.get("procedure_prose", "")
        if prose:
            h2 = self.doc.add_heading(level=2)
            r2 = h2.add_run("Procedure")
            r2.font.name = self.font_name
            r2.font.color.rgb = hex_to_rgb(self.secondary_hex)

            self._render_structured_prose(prose)
            self.doc.add_paragraph()

        # Element Descriptions Table
        fields = content_data.get("field_descriptions", [])
        if fields:
            h3 = self.doc.add_heading(level=2)
            r3 = h3.add_run("Interface Element Dictionary")
            r3.font.name = self.font_name
            r3.font.color.rgb = hex_to_rgb(self.secondary_hex)
            h3.paragraph_format.space_after = Pt(8)

            table = self.doc.add_table(rows=1, cols=2)
            table.autofit = False
            self.set_table_borders(table, "E2E8F0")

            table.columns[0].width = Inches(2.2)
            table.columns[1].width = Inches(4.3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'UI Element / Field'
            hdr_cells[1].text = 'Purpose & Functional Description'

            for idx in range(2):
                self.set_cell_background(hdr_cells[idx], self.primary_hex)
                self.set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
                p = hdr_cells[idx].paragraphs[0]
                r = p.runs[0]
                r.font.name = self.font_name
                r.font.size = Pt(9.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

            for r_idx, field in enumerate(fields):
                row_cells = table.add_row().cells
                # Strip section prefix if present (e.g. "Login Form -> Username")
                raw_name = field.get("field_name", "Unnamed Element")
                display_name = raw_name.split(" -> ")[-1] if " -> " in raw_name else raw_name
                row_cells[0].text = display_name
                row_cells[1].text = field.get("description", "No description provided.")

                fill_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c in row_cells:
                    self.set_cell_background(c, fill_color)
                    self.set_cell_margins(c, top=100, bottom=100, left=150, right=150)
                    p = c.paragraphs[0]
                    if p.runs:
                        p.runs[0].font.name = self.font_name
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.color.rgb = hex_to_rgb(self.text_hex)

            p_end = self.doc.add_paragraph()
            p_end.paragraph_format.space_before = Pt(8)

        self.doc.add_page_break()

    def _render_structured_prose(self, prose: str):
        """Parse the LLM-generated structured prose into properly styled document sections."""
        import re

        # Known section markers from the new prompt format
        SECTIONS = ["PURPOSE:", "PREREQUISITES:", "STEPS:", "EXPECTED OUTCOME:", "NOTES:"]
        SECTION_LABELS = {
            "PURPOSE:": "Purpose",
            "PREREQUISITES:": "Prerequisites",
            "STEPS:": "Steps",
            "EXPECTED OUTCOME:": "Expected Outcome",
            "NOTES:": "Notes"
        }

        current_section = None
        in_prerequisites = False
        in_steps = False

        for raw_line in prose.split("\n"):
            line = raw_line.strip()
            if not line:
                continue

            # Check if this line opens a new section
            matched_section = None
            for marker in SECTIONS:
                if line.upper().startswith(marker):
                    matched_section = marker
                    break

            if matched_section:
                current_section = matched_section
                in_prerequisites = (matched_section == "PREREQUISITES:")
                in_steps = (matched_section == "STEPS:")
                label = SECTION_LABELS[matched_section]

                if matched_section == "NOTES:":
                    # Notes go into a callout — grab remaining text on same line if any
                    remaining = line[len(matched_section):].strip()
                    if remaining:
                        self._render_callout_line(remaining)
                    continue

                # Sub-section heading
                h = self.doc.add_heading(level=3)
                r = h.add_run(label)
                r.font.name = self.font_name
                r.font.size = Pt(11)
                r.font.color.rgb = hex_to_rgb(self.primary_hex)
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after = Pt(4)

                # Inline text after section header (e.g. "PURPOSE: Manage users")
                inline = line[len(matched_section):].strip()
                if inline:
                    self._add_body_paragraph(inline)
                continue

            # Line content within a section
            if current_section == "NOTES:":
                self._render_callout_line(line)
            elif in_prerequisites and line.startswith("-"):
                p = self.doc.add_paragraph(line.lstrip("- "), style="List Bullet")
                if p.runs:
                    p.runs[0].font.name = self.font_name
                    p.runs[0].font.size = Pt(10)
            elif in_steps and re.match(r"^\d+\.", line):
                # Numbered step — use List Number style for proper formatting
                step_text = re.sub(r"^\d+\.\s*", "", line).strip()
                p = self.doc.add_paragraph(style="List Number")
                r = p.add_run(step_text)
                r.font.name = self.font_name
                r.font.size = Pt(10)
            elif line.lower().startswith(("note:", "warning:", "important:")):
                self._render_callout_line(line)
            else:
                self._add_body_paragraph(line)

    def _add_body_paragraph(self, text: str):
        """Adds a standard body text paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = self.font_name
        r.font.size = Pt(10.5)

    def _render_callout_line(self, line: str):
        """Detect Note/Warning/Important prefix and render as a callout box."""
        line_lower = line.lower()
        c_type, clean = "note", line
        if line_lower.startswith("warning:"):
            c_type = "warning"
            clean = line.split(":", 1)[1].strip()
        elif line_lower.startswith("important:"):
            c_type = "important"
            clean = line.split(":", 1)[1].strip()
        elif line_lower.startswith("note:"):
            c_type = "note"
            clean = line.split(":", 1)[1].strip()
        if clean:
            self.add_callout(clean, callout_type=c_type)

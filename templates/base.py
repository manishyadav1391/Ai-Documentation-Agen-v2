from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentBuilder(ABC):
    """
    FR-36: Base interface for pluggable client document templates.
    Any new client format (e.g., GSTAT, AIASL) must implement this class.
    """
    
    def __init__(self, style_config: Dict[str, Any] = None):
        self.doc = Document()
        self.style_config = style_config or {}
        self._setup_styles()

    @abstractmethod
    def _setup_styles(self):
        """Setup custom client fonts, colors, margins."""
        pass

    @abstractmethod
    def add_cover_page(self, title: str):
        """Generates the client-specific cover page."""
        pass

    @abstractmethod
    def add_revision_history(self):
        """Generates the document version control table."""
        pass

    @abstractmethod
    def add_toc_placeholder(self):
        """Sets up the Table of Contents structure."""
        pass

    @abstractmethod
    def add_screen_section(self, screen_index: int, image_path: Path, content_data: Dict[str, Any]):
        """
        Formats a single screen's annotated image, procedure prose, 
        and field description table.
        """
        pass

    def save(self, output_path: Path):
        """Saves the final .docx file to disk."""
        self.doc.save(output_path)

    # --- Concrete Primitives / Formatting Helpers ---

    def set_margins(self, top: float = 1.0, bottom: float = 1.0, left: float = 1.0, right: float = 1.0):
        """Set standard margins for all sections in Inches."""
        for section in self.doc.sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)

    def add_page_number_to_run(self, run):
        """Appends the PAGE field XML to a run."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def add_total_pages_to_run(self, run):
        """Appends the NUMPAGES field XML to a run."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "NUMPAGES"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def setup_header_footer(self, logo_path: Path, doc_title: str, font_name: str = 'Arial'):
        """Creates a professional logo header (left) and document title (right) + footer numbering.
        Safe to call multiple times — clears existing header content before writing.
        Only applies to non-first-page sections (cover page is excluded via different_first_page_header_footer).
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        for section in self.doc.sections:
            section.different_first_page_header_footer = True

            # ── HEADER ────────────────────────────────────────────────────────
            header = section.header
            header.is_linked_to_previous = False

            # Fully clear ALL existing paragraphs and tables in the header XML
            hdr_element = header._element
            for child in list(hdr_element):
                hdr_element.remove(child)

            # Re-create the header body paragraph (Word requires at least one)
            # Use a table for logo-left / title-right layout
            hdr_p = OxmlElement('w:p')
            hdr_element.append(hdr_p)

            table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            table.autofit = False

            # Remove all table borders
            tblPr = table._tbl.tblPr
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tblBorders.append(border)
            tblPr.append(tblBorders)

            table.columns[0].width = Inches(3.25)
            table.columns[1].width = Inches(3.25)

            # Cell 0: Logo (left-aligned)
            cell_logo = table.cell(0, 0)
            p_logo = cell_logo.paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if logo_path and Path(logo_path).exists():
                try:
                    p_logo.add_run().add_picture(str(logo_path), height=Inches(0.38))
                except Exception:
                    pass

            # Cell 1: Document title (right-aligned)
            cell_title = table.cell(0, 1)
            p_title = cell_title.paragraphs[0]
            p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_title = p_title.add_run(doc_title)
            run_title.font.name = font_name
            run_title.font.size = Pt(8.5)
            run_title.font.italic = True
            run_title.font.color.rgb = RGBColor(100, 100, 100)

            # Thin separator line below header table
            sep_p = header.add_paragraph()
            self.add_bottom_border_to_paragraph(sep_p, color_hex="CCCCCC", sz=4)
            sep_p.paragraph_format.space_before = Pt(2)
            sep_p.paragraph_format.space_after = Pt(0)

            # ── FOOTER ────────────────────────────────────────────────────────
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear footer content
            ftr_element = footer._element
            for child in list(ftr_element):
                ftr_element.remove(child)

            ftr_p_xml = OxmlElement('w:p')
            ftr_element.append(ftr_p_xml)

            p_foot = footer.paragraphs[0]

            # Left: Company name
            p_foot.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Use a two-column table for footer: company left, page number right
            ftr_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
            ftr_table.autofit = False

            ftr_tblPr = ftr_table._tbl.tblPr
            ftr_borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                b = OxmlElement(f'w:{border_name}')
                b.set(qn('w:val'), 'none')
                ftr_borders.append(b)
            ftr_tblPr.append(ftr_borders)

            ftr_table.columns[0].width = Inches(3.25)
            ftr_table.columns[1].width = Inches(3.25)

            company = self.style_config.get("company_name", "")
            p_left = ftr_table.cell(0, 0).paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r_company = p_left.add_run(company)
            r_company.font.name = font_name
            r_company.font.size = Pt(8)
            r_company.font.color.rgb = RGBColor(130, 130, 130)

            p_right = ftr_table.cell(0, 1).paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_pre = p_right.add_run("Page ")
            r_pre.font.name = font_name
            r_pre.font.size = Pt(8)
            r_pre.font.color.rgb = RGBColor(100, 100, 100)
            self.add_page_number_to_run(p_right.add_run())
            r_mid = p_right.add_run(" of ")
            r_mid.font.name = font_name
            r_mid.font.size = Pt(8)
            r_mid.font.color.rgb = RGBColor(100, 100, 100)
            self.add_total_pages_to_run(p_right.add_run())


    def set_cell_background(self, cell, fill_hex: str):
        """Set cell background shading XML."""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_table_borders(self, table, border_color_hex: str = "CCCCCC"):
        """Sets custom borders on all cells in a table."""
        tblPr = table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') # border width (sz of 4 is 1/2 pt)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color_hex)
            tblBorders.append(border)
        tblPr.append(tblBorders)

    def set_cell_margins(self, cell, top: int = 120, bottom: int = 120, left: int = 180, right: int = 180):
        """Set cell padding (in dxa: 1 pt = 20 dxa). Default is top/bottom=6pt, left/right=9pt."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_cell_left_border(self, cell, color_hex: str, size_sz: int = 36):
        """Set a thick left border on a cell, clearing other borders (useful for callouts)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        # Left border: custom color and size (size_sz: 36 is 4.5 pt)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(size_sz))
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), color_hex)
        tcBorders.append(left)
        
        # Clear top, bottom, right borders
        for border_name in ['top', 'bottom', 'right']:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
            
        tcPr.append(tcBorders)

    def add_toc_field(self, paragraph):
        """Adds a native Microsoft Word TOC field to a paragraph."""
        p_element = paragraph._p
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        p_element.append(fldChar1)
        p_element.append(instrText)
        p_element.append(fldChar2)
        
        r = OxmlElement('w:r')
        rText = OxmlElement('w:t')
        rText.text = "[Right-click here and select 'Update Field' to generate Table of Contents]"
        r.append(rText)
        p_element.append(r)
        p_element.append(fldChar3)

    def add_bottom_border_to_paragraph(self, paragraph, color_hex: str = "CCCCCC", sz: int = 6):
        """Appends a bottom border line to a paragraph, functioning as a clean divider."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(sz))
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)